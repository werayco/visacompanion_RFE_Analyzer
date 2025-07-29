from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

class utils:
    @staticmethod
    def CriterionDocCreator(result: dict) -> str:
        doc = Document()
        
        title = doc.add_heading("EB-1A Petition Assessment Summary", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        section_num = 1

        recognizable_components = result.get("recognizable_components", {})
        background = recognizable_components.get("personal_background", "")
        expert_letters = recognizable_components.get("expert_recommendation_letters", [])
        media_and_judging = recognizable_components.get("media_and_judging_roles", [])
        red_flags = recognizable_components.get("red_flags", {})
        evidence = recognizable_components.get("evidence", {})

        doc.add_heading(f"{section_num}. Personal Background", level=2)
        if background:
            doc.add_paragraph(background)
        else:
            doc.add_paragraph("No personal background information provided.")
        section_num += 1

        doc.add_heading(f"{section_num}. Expert Recommendation Letters", level=2)
        if expert_letters:
            for i, letter in enumerate(expert_letters, 1):
                doc.add_paragraph(f"{i}. {letter}")
        else:
            doc.add_paragraph("No expert recommendation letters documented.")
        section_num += 1

        doc.add_heading(f"{section_num}. Media Coverage and Judging Activities", level=2)
        if media_and_judging:
            for role in media_and_judging:
                doc.add_paragraph(f" {role}")
        else:
            doc.add_paragraph("No media coverage or judging activities documented.")
        section_num += 1

        doc.add_heading(f"{section_num}. Red Flags Analysis", level=2)
        if red_flags:
            for k, v in red_flags.items():
                title_text = k.replace("_", " ").title()
                doc.add_paragraph(f" {title_text}: {v}")
        else:
            doc.add_paragraph("No red flags analysis provided.")
        section_num += 1
        internationalAward = result.get("internationalAward", {})
        doc.add_heading(f"{section_num}. International Award Evaluation", level=2)
        
        majorAwards = internationalAward.get("has_major_award", False)
        has_award_text = "✓ Yes, a major internationally recognized award is present." if majorAwards else "✗ No, a qualifying major internationally recognized award was found."
        
        award_para = doc.add_paragraph()
        award_para.add_run("Has Major Award: ").bold = True
        award_para.add_run(has_award_text)

        award_names = internationalAward.get("award_names", [])
        if award_names:
            awards_para = doc.add_paragraph()
            awards_para.add_run("Award(s) Mentioned: ").bold = True
            for name in award_names:
                doc.add_paragraph(f" {name}", style="List Bullet")

        justification = internationalAward.get("justification", "")
        if justification:
            just_para = doc.add_paragraph()
            just_para.add_run("Justification: ").bold = True
            doc.add_paragraph(justification)

        missing_elements = internationalAward.get("missing_elements", [])
        if missing_elements:
            missing_para = doc.add_paragraph()
            missing_para.add_run("Missing Elements: ").bold = True
            for item in missing_elements:
                doc.add_paragraph(f" {item}", style="List Bullet")

        suggested_supporting_evidence = internationalAward.get("suggested_supporting_evidence", [])
        if suggested_supporting_evidence:
            suggest_para = doc.add_paragraph()
            suggest_para.add_run("Suggested Supporting Evidence: ").bold = True
            for item in suggested_supporting_evidence:
                doc.add_paragraph(f" {item}", style="List Bullet")
        
        section_num += 1

        summary = result.get("summary", {})
        doc.add_heading(f"{section_num}. Executive Summary", level=2)
        
        summary_items = [
            ("Qualification Path", summary.get('qualification_path', 'N/A')),
            ("Overall Status", summary.get('status', 'N/A')),
            ("Criteria Met", f"{len(summary.get('met_criteria', []))} out of {summary.get('criteria_needed', 3)} required"),
            ("Total Criteria Evaluated", summary.get('total_criteria_checked', 'N/A')),
            ("Met Criteria", ', '.join([c.replace('Result', '').replace('Criterion', 'Criterion ') for c in summary.get('met_criteria', [])]))
        ]
        
        for label, value in summary_items:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))
        
        section_num += 1

        doc.add_heading(f"{section_num}. Detailed Criterion Evaluation", level=2)

        criterionTitles = {
            "Criterion1Result": "Criterion 1: Lesser Nationally/Internationally Recognized Awards",
            "Criterion2Result": "Criterion 2: Membership in Distinguished Associations", 
            "Criterion3Result": "Criterion 3: Published Material About the Beneficiary",
            "Criterion4Result": "Criterion 4: Judging Others' Work",
            "Criterion5Result": "Criterion 5: Original Contributions of Major Significance",
            "Criterion6Result": "Criterion 6: Scholarly Articles Authored",
            "Criterion7Result": "Criterion 7: Artistic Work Display/Exhibition",
            "Criterion8Result": "Criterion 8: Leading/Critical Role in Distinguished Organizations",
            "Criterion9Result": "Criterion 9: High Salary/Remuneration",
            "Criterion10Result": "Criterion 10: Commercial Success in Performing Arts",
        }

        met_criteria = set(summary.get('met_criteria', []))

        for i in range(1, 11):
            criterion_key = f"Criterion{i}Result"
            criterion_data = result.get(criterion_key, {})
            
            title = criterionTitles.get(criterion_key, f"Criterion {i}")
            is_met = criterion_key in met_criteria
            status_indicator = " ✓ MET" if is_met else " ✗ NOT MET"
            
            heading = doc.add_heading(title + status_indicator, level=3)
            if is_met:
                heading.style.font.color.rgb = None

            if not criterion_data:
                doc.add_paragraph(" No data provided for this criterion.")
                continue

            if criterion_data.get("status") == "couldn't parse the content":
                doc.add_paragraph(" Could not parse the content for this criterion.")
                continue

            meets_criterion = criterion_data.get(f'meets_criterion_{i}')
            if meets_criterion is not None:
                status_para = doc.add_paragraph()
                status_para.add_run("Status: ").bold = True
                status_text = "MEETS CRITERION" if meets_criterion else "DOES NOT MEET CRITERION"
                status_para.add_run(status_text)

            reasoning = criterion_data.get("reasoning") or criterion_data.get("overall_reasoning", "")
            if reasoning:
                reason_para = doc.add_paragraph()
                reason_para.add_run("Analysis: ").bold = True
                doc.add_paragraph(reasoning)

            utils._add_criterion_details(doc, criterion_data, i)

        output_path = "Criterion_EB1A_Assessment_Summary.docx"
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def _add_criterion_details(doc, criterion_data, criterion_num):
        
        if "award_names" in criterion_data and criterion_data["award_names"]:
            doc.add_paragraph().add_run("Awards Identified:").bold = True
            for award in criterion_data["award_names"]:
                doc.add_paragraph(f" {award}", style="List Bullet")

        if "association_names" in criterion_data and criterion_data["association_names"]:
            doc.add_paragraph().add_run("Associations:").bold = True
            for assoc in criterion_data["association_names"]:
                doc.add_paragraph(f" {assoc}", style="List Bullet")

        if "published_materials" in criterion_data and criterion_data["published_materials"]:
            doc.add_paragraph().add_run("Published Materials:").bold = True
            for material in criterion_data["published_materials"]:
                doc.add_paragraph(f" {material}", style="List Bullet")

        if "judging_instances" in criterion_data and criterion_data["judging_instances"]:
            doc.add_paragraph().add_run("Judging Activities:").bold = True
            for instance in criterion_data["judging_instances"]:
                if isinstance(instance, dict):
                    event = instance.get("event_name_or_platform", "Unknown Event")
                    role = instance.get("role", "Unknown Role")
                    date = instance.get("date_or_period", "Unknown Date")
                    doc.add_paragraph(f" {event} - {role} ({date})", style="List Bullet")
                else:
                    doc.add_paragraph(f" {instance}", style="List Bullet")

        if "contributions" in criterion_data and criterion_data["contributions"]:
            doc.add_paragraph().add_run("Major Contributions:").bold = True
            for contrib in criterion_data["contributions"]:
                title = contrib.get("title_or_description", "Untitled Contribution")
                impact = contrib.get("impact_summary", "")
                significance = contrib.get("is_major_significance", False)
                
                contrib_para = doc.add_paragraph(f" {title}", style="List Bullet")
                if significance:
                    contrib_para.add_run(" [MAJOR SIGNIFICANCE]").bold = True
                
                if impact:
                    doc.add_paragraph(f"  Impact: {impact}")

        if "articles" in criterion_data and criterion_data["articles"]:
            doc.add_paragraph().add_run("Authored Articles:").bold = True
            for article in criterion_data["articles"]:
                if isinstance(article, dict):
                    title = article.get("title", "Untitled")
                    publication = article.get("publication", "Unknown Publication")
                    doc.add_paragraph(f" \"{title}\" in {publication}", style="List Bullet")
                else:
                    doc.add_paragraph(f" {article}", style="List Bullet")

        if "exhibitions" in criterion_data and criterion_data["exhibitions"]:
            doc.add_paragraph().add_run("Exhibitions/Displays:").bold = True
            for exhibition in criterion_data["exhibitions"]:
                doc.add_paragraph(f" {exhibition}", style="List Bullet")

        if "leading_roles" in criterion_data and criterion_data["leading_roles"]:
            doc.add_paragraph().add_run("Leading/Critical Roles:").bold = True
            for role in criterion_data["leading_roles"]:
                doc.add_paragraph(f" {role}", style="List Bullet")

        if "salary_records" in criterion_data and criterion_data["salary_records"]:
            doc.add_paragraph().add_run("Salary Information:").bold = True
            for record in criterion_data["salary_records"]:
                if isinstance(record, dict):
                    year = record.get("year", "Unknown Period")
                    amount = record.get("salary_amount", "Unknown Amount")
                    is_high = record.get("is_significantly_high", False)
                    high_indicator = " [SIGNIFICANTLY HIGH]" if is_high else ""
                    doc.add_paragraph(f" {year}: {amount}{high_indicator}", style="List Bullet")
                else:
                    doc.add_paragraph(f" {record}", style="List Bullet")

        if "works" in criterion_data and criterion_data["works"]:
            doc.add_paragraph().add_run("Commercial Successes:").bold = True
            for work in criterion_data["works"]:
                doc.add_paragraph(f" {work}", style="List Bullet")

        if "issues_detected" in criterion_data and criterion_data["issues_detected"]:
            doc.add_paragraph().add_run("Issues Identified:").bold = True
            for issue in criterion_data["issues_detected"]:
                doc.add_paragraph(f" {issue}", style="List Bullet")

        if "missing_elements" in criterion_data and criterion_data["missing_elements"]:
            doc.add_paragraph().add_run("Missing Elements:").bold = True
            for element in criterion_data["missing_elements"]:
                doc.add_paragraph(f" {element}", style="List Bullet")

        if "suggested_supporting_evidence" in criterion_data and criterion_data["suggested_supporting_evidence"]:
            doc.add_paragraph().add_run("Suggested Supporting Evidence:").bold = True
            for evidence in criterion_data["suggested_supporting_evidence"]:
                doc.add_paragraph(f" {evidence}", style="List Bullet")

    @staticmethod
    def jsonSaver(jsonObject: dict):
        with open("result.json", "w", encoding="utf-8") as result:
            json.dump(jsonObject, result, ensure_ascii=False, indent=2)

    @staticmethod
    def JsonLoader(jsonPath: str) -> dict:
        with open(jsonPath, "r") as result:
            data = json.load(result)
            return data

# json_data = utils.JsonLoader(r"result.json")
# output_file = utils.CriterionDocCreator(json_data)
