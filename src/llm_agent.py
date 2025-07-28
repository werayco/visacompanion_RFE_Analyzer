from docx import Document
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from langchain_groq import ChatGroq
import re
import os
import json
from typing import Dict
from langchain_anthropic import ChatAnthropic
import time
from dotenv import load_dotenv

load_dotenv()

groq_api_key = os.getenv("GROQ_API_KEY")
claude_api_key = os.getenv("CLAUDE_API_KEY")
agent_temp = os.getenv("AGENT_TEMP")

if not groq_api_key and not claude_api_key or agent_temp:
    raise ValueError("Kindly make sure your environmental keys are present in the .env file.")


def llm(question: str, model: str = "deepseek-r1-distill-llama-70b") -> str:
    """groq llm"""
    try:
        chat = ChatGroq(model=model, api_key=groq_api_key, temperature=0.1)
        response = chat.invoke(question)
        if response:
            return response.content
    except Exception as e:
        return f"An error occurred: {e}"


def llm2(question: str):
    """claude's llm"""
    try:
        llm2 = ChatAnthropic(
            model="claude-3-opus-20240229",
            temperature=agent_temp,
            api_key=claude_api_key,
        )
        response = llm2.invoke(question)
        if response:
            return response.content
    except Exception as e:
        return f"An Error Occured, {e}"


class RFEBOT:
    @staticmethod
    def reponseParser(text):
        text = text.strip()

        def clean_json_string(match):
            return match.group(0).replace("\n", " ").replace("\r", "")

        try:
            pattern = r"```(?:json)?\s*(\{.*?\})\s*```"
            match = re.search(pattern, text, re.DOTALL)
            if match:
                json_text = match.group(1).strip()
                json_text = re.sub(r'"[^"]*"', clean_json_string, json_text)
                return json.loads(json_text)
        except Exception as e:
            pass
        try:
            if text.startswith("{") and text.endswith("}"):
                cleaned_text = re.sub(r'"[^"]*"', clean_json_string, text)
                return json.loads(cleaned_text)
        except Exception as e:
            pass
        try:
            pattern = r"(\{.*\})"
            match = re.search(pattern, text, re.DOTALL)
            if match:
                json_text = match.group(1).strip()
                json_text = re.sub(r'"[^"]*"', clean_json_string, json_text)
                return json.loads(json_text)
        except Exception as e:
            pass
        return {"status": "couldn't parse the content"}
    @staticmethod
    def Agentprompt(petitionContent: str) -> Dict:
        prompt1: str = f"""You are a USCIS Officer assigned to critically review an EB-1A petition. Your task is to extract the following components from the petition content:

        1. Personal Background  
        Summarize biographical details, including education, work history, and other relevant personal information that provides context for the petition.

        2. Expert Recommendation Letters  
        Identify and extract any recommendation letters written by domain experts attesting to the petitioner's extraordinary ability, achievements, or influence in the field.

        3. Media Coverage and Judging Roles  
        Highlight any media features, interviews, or evidence of judging/peer review activities.

        # Identify Potential Red Flags  
        Carefully evaluate the petition for any signs that could weaken the credibility or strength of the evidence provided. Flag any of the following:

        4. Repetitive or Boilerplate Language  
        - Identical or overly similar wording across multiple recommendation letters  
        - Copy-pasted templates with minimal personalization or field-specific insight

        5. Generic or Vague Contribution Claims  
        - Statements like "significant impact" or "major innovation" without detailed explanation, technical depth, or third-party validation  
        - Lack of specificity about the petitioner's unique role or achievements

        6. Insufficient Supporting Evidence  
        - Claims made without verifiable documentation (e.g., links, data, press mentions, awards, citations)  
        - No corroborating metrics (e.g., usage stats, revenue, adoption, expert testimonials)

        7. Unverified or Unremarkable Salary Information  
        - Salary claims without comparative benchmarks or third-party validation (e.g., industry averages, country-specific data)  
        - High earnings not contextualized within the petitioner's field or region

        8. Inconsistent or Broad Field of Expertise  
        - Evidence presented across unrelated or loosely connected domains (e.g., AI + graphic design + biology)  
        - Lack of a clear, focused field that aligns with the claimed extraordinary ability

        Below is the Petition Content:  
        <petitionContent>  
        {petitionContent.strip()}  
        </petitionContent>

        Return the results in a JSON object with the following structure:
        {{
            "personal_background": "...",
            "expert_recommendation_letters": ["...", "..."],
            "media_and_judging_roles": ["..."],
            "red_flags": {{
                "repetitive_language": "...",
                "generic_claims": "...",
                "weak_substantiation": "...",
                "unsupported_salary": "...",
                "inconsistent_field": "..."
            }}
        }}
        """
        prompt2: str = f"""You are a USCIS Officer assigned to critically review an EB-1A petition. Your task is to extract evidence supporting each EB-1A criterion from the petition content.

        For each applicable criterion below, extract and organize corresponding documentation:
        (i) Receipt of lesser nationally or internationally recognized prizes or awards for excellence in the field.
        (ii) Membership in associations that require outstanding achievements of their members.
        (iii) Published material about the alien in professional or major trade publications or other major media.
        (iv) Participation as a judge of the work of others in the same or allied field.
        (v) Original contributions of major significance in the field.
        (vi) Authorship of scholarly articles in professional or major trade publications or major media.
        (vii) Display of the alien's work in artistic exhibitions or showcases.
        (viii) Leading or critical roles for distinguished organizations or establishments.
        (ix) Evidence of commanding a high salary or significantly high remuneration relative to others in the field.
        (x) Commercial success in the performing arts, such as box office receipts, record/cassette/CD/video sales.

        Extract all the sections that involve each criterion from the original petition content. The extracted content for each criterion must be verbatim and unchanged from the source. If a criterion is not supported by any content, return an empty array for that criterion.

        Below is the Petition Content:
        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Return the results in a JSON object with the following structure:
        {{
            "evidence": {{
                "lesser_awards": ["..."],
                "association_membership": ["..."],
                "published_material": ["..."],
                "judging_work_of_others": ["..."],
                "original_contributions": ["..."],
                "scholarly_articles": ["..."],
                "artistic_display": ["..."],
                "leading_or_critical_role": ["..."],
                "high_salary_or_remuneration": ["..."],
                "commercial_success_in_arts": ["..."]
            }}
        }}"""

        llm2_response1 = llm2(prompt1)
        llm2_response2 = llm2(prompt2)

        response1 = RFEBOT.reponseParser(llm2_response1)
        response2 = RFEBOT.reponseParser(llm2_response2)

        combined_response = {
            "personal_background": response1.get("personal_background", ""),
            "expert_recommendation_letters": response1.get(
                "expert_recommendation_letters", []
            ),
            "evidence": response2.get("evidence", {}),
            "media_and_judging_roles": response1.get("media_and_judging_roles", []),
            "red_flags": response1.get("red_flags")
        }

        return combined_response

    @staticmethod
    def AgentpromptInternationalAward(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies the "one-time achievement" criterion—specifically, whether they have received a **major, internationally recognized award**.

        Below is the Petition Content:

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Instructions:

        1. Check if the document contains any evidence of the beneficiary receiving a **major, internationally recognized award** (e.g., Nobel Prize, Olympic Medal, Pulitzer Prize).
        2. If such an award is found, set `has_major_award` to `true` and list the award(s) under `award_names`.
        3. If the award(s) mentioned do **not appear to qualify**, set `has_major_award` to `false` and explain why in `justification`.
        4. Refer to the following evaluation factors to determine whether the award qualifies:
        - Is the award internationally recognized in the field?
        - Does it attract worldwide competition?
        - Has it been reported internationally in top media?
        - Is it a familiar name to the general public?
        - Does it include a large cash prize?
        5. If any of these elements are **missing or insufficient**, list them under `missing_elements`.
        6. Under `suggested_supporting_evidence`, propose what kind of documentation could strengthen the petitioner's claim that the award qualifies (based on the RFE guidance).

        Note: Use concise, formal language in your response.

        ALWAYS Return the results in a JSON object with the following structure:

        {{
        "has_major_award": true | false,
        "award_names": ["..."],
        "justification": "...",
        "missing_elements": ["..."], 
        "suggested_supporting_evidence": ["..."]
        }}
        """
        llm2_response = llm2(prompt)
        response = RFEBOT.reponseParser(llm2_response)
        return response

    @staticmethod
    def AgenticCriterionPrompt1(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer evaluating EB-1A Criterion 1: "Documentation of the alien's receipt of lesser nationally or internationally recognized prizes or awards for excellence in the field of endeavor."

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Task:
        1. Extract award names into `award_names` field
        2. Set `meets_criterion_1` to `true` only if awards have strong evidence of being:
        - Nationally/internationally recognized
        - Granted for excellence in beneficiary's specific field

        3. Address these RFE concerns:
        - Award scope (local/regional vs national/international)
        - Purpose (excellence vs participation/attendance)
        - Field relevance
        - Academic-only status
        - Documentary evidence gaps:
            * Selection criteria * Award significance/prestige
            * Granting organization reputation * Geographic eligibility scope
            * Annual recipient numbers * Notable past winners
            * Media coverage/announcements

        4. Provide analysis in JSON format with decision reasoning, detected issues, missing elements, and suggested supporting evidence.

        Return only JSON:
        {{
        "meets_criterion_1": true | false,
        "award_names": ["..."],
        "reasoning": "...",
        "issues_detected": ["..."],
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}"""
        llm2_response = llm2(prompt)
        response = RFEBOT.reponseParser(llm2_response)
        return response

    @staticmethod
    def AgenticCriterionPrompt2(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 2**: “Documentation of the beneficiary's membership in associations in the field for which classification is sought. The association must require outstanding achievements of their members, as judged by recognized national or international experts in their disciplines or fields.”

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>


        Instructions:

        1. Identify all associations the beneficiary claims membership in and list them under `association_names`.

        2. Set `meets_criterion_2` to `true` **only if** the following are established:
        - The association is clearly related to the petitioner's field.
        - The association **requires outstanding achievements** for the level of membership held.
        - The selection is judged or reviewed by **recognized national or international experts** in the field.

        3. Carefully evaluate for RFE indicators. If applicable, list problems in `issues_detected`. Common issues include:
        - No proof that the association is **related to the field**.
        - Membership appears to be **open or general** rather than achievement-based.
        - No evidence that reviewers are **recognized experts**.
        - No **bylaws or membership criteria** included.
        - Only student-level or early-career membership provided.

        4. Use `missing_elements` to highlight gaps in evidence. These may include:
        - Lack of association bylaws.
        - No information on reviewer qualifications.
        - No criteria showing outstanding achievements are required.
        - Lack of clarity about the beneficiary's **level of membership** (e.g., Fellow vs. General).

        5. In `suggested_supporting_evidence`, recommend helpful documentation, such as:
        - The relevant section of the association's constitution/bylaws showing membership requirements.
        - Documents proving that selection is based on outstanding achievements.
        - Info establishing that reviewers are nationally or internationally recognized experts.
        - Evidence that association membership is limited and prestigious.

        Return only the JSON object. Do not add explanation outside the object.
        Analyze the content and return your findings in the following JSON format:

        {{
        "meets_criterion_2": true | false,
        "association_names": ["..."],
        "reasoning": "...",
        "issues_detected": ["..."],
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        response = RFEBOT.reponseParser(llm2(prompt))
        return response  # this is the reponse for second criterion

    @staticmethod
    def AgenticCriterionPrompt3(petitionContent: str) -> Dict:
        prompt = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 3**: "Published material about the beneficiary in professional or major trade publications or other major media." The materials must relate to the beneficiary’s work in the field for which classification is sought.

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Instructions:

        1. For each piece of published material, extract key information including:
        - Title, author, date
        - Media type (print, online, broadcast)
        - Publication name
        - Circulation or reach (if provided)
        - Intended audience (professional, general, etc.)

        2. Determine:
        - Is the content **about the beneficiary**? (`about_beneficiary`)
        - Is it specifically about **the beneficiary's work in the field**? (`about_beneficiarys_work`)

        3. Set `meets_criterion_3` to `true` only if:

        - At least one published material meets both criteria above.
        - The publication is a **professional or major trade publication** or **major media outlet**.
        - The content focuses on the **beneficiary and their specific work** (not just the company or a product, and not a mere mention or citation).

        4. If the content fails on any of the following, include them in `issues` per article:
        - Publication is not recognized as major media or trade publication.
        - Article is not actually about the beneficiary or their work.
        - It's a marketing or promotional piece, not journalistic content.
        - Citations or footnotes without evaluative content.
        - No proof of reach, relevance, or media credibility.

        5. In `missing_elements`, list key supporting details that are absent (e.g., "no circulation info," "missing publication audience").

        6. In `suggested_supporting_evidence`, recommend documents the petitioner could include to strengthen this criterion, such as:
        - Screenshots or scans of articles showing full publication data
        - Media kits or audience reports of the publication
        - Articles with in-depth focus on the beneficiary's specific achievements
        - Evidence of independent reporting (not self-promotional)

        You must **ALWAYS** Return the result in the following JSON format:

        {{
        "meets_criterion_3": true | false,
        "published_materials": [
            {{
            "title": "...",
            "author": "...",
            "date": "...",
            "media_type": "print | online | broadcast",
            "publication_name": "...",
            "circulation_or_reach": "...",
            "intended_audience": "...",
            "about_beneficiary": true | false,
            "about_beneficiarys_work": true | false,
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        response = RFEBOT.reponseParser(llm2(prompt))
        return response  # this is the reponse for criterion 3

    @staticmethod
    def AgenticCriterionPrompt4(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 4**: "Evidence of the beneficiary's participation, either individually or on a panel, as a judge of the work of others in the same or an allied field of specialization for which classification is sought."

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Instructions:

        1. For each judging activity submitted, extract:
        - Event name or organization/platform (e.g. journal, conference, grant review)
        - Role (e.g. manuscript reviewer, panel judge, grant evaluator)
        - Date or time period of involvement
        - Type of work evaluated (e.g. academic papers, projects, startup proposals)
        - Field in which the work was judged

        2. Determine if the judging activity was in the **same** or an **allied** field to the beneficiary’s claimed field of expertise.
        - If yes, set `same_or_allied_field: true`
        - If no, or unclear, set `same_or_allied_field: false` and include justification in `issues`

        3. Set `meets_criterion_4: true` only if at least one clear, documented instance shows the beneficiary judged work of others **in the same or allied field**.

        4. Include in `issues` if:
        - The judged field appears unrelated to the beneficiary's claimed field.
        - There is no evidence of actual judging (e.g. invitation without confirmation).
        - The event or organization is not credible or lacks detail.

        5. List `missing_elements` such as:
        - No link between judging field and claimed field
        - No proof of participation
        - No date or role specified

        6. Recommend `suggested_supporting_evidence` such as:
        - Copies or screenshots of review requests and confirmations
        - Proof of judging (e.g. thank-you letters, portal screenshots)
        - Descriptions or links showing the relevance of the judged material to the beneficiary’s field
        - Statements or credentials showing how the reviewed field is allied

        You must **ALWAYS** Return the result in the following JSON format. Do not include commentary outside the structure.

        {{
        "meets_criterion_4": true | false,
        "judging_instances": [
            {{
            "event_name_or_platform": "...",
            "role": "...",
            "date_or_period": "...",
            "type_of_work_reviewed": "...",
            "field_of_judging": "...",
            "same_or_allied_field": true | false,
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        response = RFEBOT.reponseParser(llm2(prompt))
        return response  # this is a prompt for the criterion 4

    @staticmethod
    def AgenticCriterionPrompt5(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 5**: "Evidence of the beneficiary's original scientific, scholarly, artistic, athletic, or business-related contributions of major significance in the field."

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Instructions:

        1. Extract all described contributions from the submission, including publications, innovations, business ventures, designs, software, athletic feats, or scholarly theories.

        2. For each, assess:
        - Whether the work is **original** (i.e., novel and not a replication)
        - Whether the contribution is of **major significance** (i.e., widely cited, adopted, licensed, or used across the field)
        - Whether experts or credible entities in the field recognize the importance of the work

        3. Set `meets_criterion_5: true` only if at least one contribution is clearly both original and of major significance to the field of endeavor.

        4. In `evidence_summary`, include what was submitted (e.g., journal articles, patents, letters from experts, media coverage, usage contracts, licenses).

        5. In `recognized_by_field`, mark `true` if there is documentary proof that experts or practitioners across the field acknowledge the importance of the work.

        6. In `impact_summary`, briefly describe how the work has influenced the field or has been implemented (e.g., adopted by organizations, licensed, cited widely).

        7. Use `issues` to flag:
        - Lack of proof of originality or uniqueness
        - Lack of implementation, citation, or real-world use
        - General praise without specific, measurable impact
        - Vague expert letters without concrete details

        8. Use `missing_elements` to list:
        - No implementation or licensing proof
        - No third-party expert commentary
        - No usage evidence
        - No citation or media reference

        9. Recommend `suggested_supporting_evidence` such as:
        - Expert letters describing how the work is both original and significant
        - Evidence of use in industry (contracts, deployment, citations, partnerships)
        - Patent citations or commercial impact data
        - Screenshots, links, or media articles referencing the innovation

        ALWAYS Return only the JSON object. Do not include any commentary or explanations outside the structure.

        {{
        "meets_criterion_5": true | false,
        "contributions": [
            {{
            "title_or_description": "...",
            "field_of_contribution": "...",
            "type_of_contribution": "...",  // scientific, scholarly, business-related, etc.
            "is_original": true | false,
            "is_major_significance": true | false,
            "evidence_summary": ["..."],
            "recognized_by_field": true | false,
            "impact_summary": "...",
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        llm2_response = llm2(prompt)
        response = RFEBOT.reponseParser(llm2_response)
        return response

    @staticmethod
    def AgenticCriterionPrompt6(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 6**: "Evidence of the beneficiary's authorship of scholarly articles in the field, in professional or major trade publications or other major media."
                
        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>


        Instructions:

        1. Extract each listed article or publication and analyze whether it qualifies under the following conditions:
        - The article must be **scholarly in nature** (i.e., deals with academic subject matter)
        - It must be **authored by the petitioner** (name should appear clearly as author)
        - It must have been published in a **professional journal, trade publication, or other major media**
        - It must preferably have undergone a **peer review or editorial process**

        2. Set `meets_criterion_6: true` only if **at least one** article clearly meets all the above conditions.

        3. For each article:
        - In `publication_type`, specify whether it's a peer-reviewed journal, trade magazine, or major media (e.g., Forbes, Nature, IEEE Transactions, Harvard Business Review).
        - In `academic_subject_matter`, mark `true` if the article focuses on a technical, business, scientific, or artistic theory or concept.
        - In `peer_reviewed`, mark `true` only if documentation of a peer review process exists or the journal is known for it.
        - In `publication_is_major`, mark `true` if the publication is widely circulated, prestigious, or highly ranked.

        4. In `evidence_summary`, include any submitted materials such as:
        - Citation indexes (Google Scholar, Scopus)
        - Screenshots or PDFs of journal covers and articles
        - Proof of peer-review process
        - Circulation metrics or media reputation

        5. Use `issues` to list any concerns, such as:
        - Lack of proof of authorship
        - Article appears in blog, newsletter, or unknown platform
        - Publication is obscure or lacks prestige
        - Article lacks academic depth

        6. Use `missing_elements` to list what is absent in the supporting documents:
        - No evidence of peer-review
        - No subject matter analysis
        - No circulation stats for publication
        - No indication it's a professional or major trade publication

        7. Recommend `suggested_supporting_evidence` such as:
        - Screenshots or links from Google Scholar showing authorship
        - Peer-review policy pages from the journal website
        - Circulation or ranking data from third-party journal databases
        - Expert testimony confirming academic nature and journal prestige

        Return only the JSON object. Do not include any commentary or explanations outside the structure.

        {{
        "meets_criterion_6": true | false,
        "articles": [
            {{
            "title": "...",
            "publication": "...",
            "publication_type": "...",  // e.g., peer-reviewed journal, trade magazine, major media
            "academic_subject_matter": true | false,
            "peer_reviewed": true | false,
            "publication_is_major": true | false,
            "evidence_summary": ["..."],
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        llm2_response = llm2(prompt)
        response = RFEBOT.reponseParser(llm2_response)
        return response

    @staticmethod
    def AgenticCriterionPrompt7(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 7**:  **"Evidence of the display of the beneficiary's work in the field at artistic exhibitions or showcases."**
                
        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>


        Instructions:

        1. Extract each exhibition or display and evaluate if it meets the following:
        - The work was **created by the petitioner** (authorship can be verified)
        - The venue was a **recognized artistic exhibition or showcase**, either physical or virtual

        2. Set `meets_criterion_7: true` only if **at least one** exhibition satisfies both:
        - Authorship is confirmed
        - Venue qualifies as a legitimate artistic exhibition or showcase

        3. For each exhibition:
        - In `venue`, include the name of the gallery, museum, conference, or online platform
        - Mark `is_artistic_showcase: true` only if the venue is a curated or recognized artistic display platform
        - Use `is_virtual` to indicate whether the exhibition was online or physical
        - Mark `authorship_verified: true` only if documentation shows the work was created by the petitioner

        4. Use `evidence_summary` to list submitted materials such as:
        - Photographs or videos of the displayed work
        - Catalogs, exhibition brochures, or gallery announcements
        - Invitations, press coverage, or promotional materials
        - Artist statements or testimony confirming authorship

        5. Use `issues` to flag any problems, such as:
        - Venue not clearly an artistic exhibition or showcase
        - Lack of evidence linking displayed work to the petitioner
        - Work was displayed but in a commercial or informal context

        6. Use `missing_elements` to list absent materials, such as:
        - No proof that the petitioner created the work
        - No documentation proving the venue is an artistic exhibition or showcase
        - No promotional or curatorial materials

        7. Use `suggested_supporting_evidence` to list ways to strengthen the petition, such as:
        - Submit sales records, creation logs, or copyright statements verifying authorship
        - Provide documentation about the venue's reputation as an artistic showcase
        - Submit promotional flyers, press articles, or social media posts by the venue

        Return the result in the following JSON format. Do not include any commentary or explanations outside the structure.
        {{
        "meets_criterion_7": true | false,
        "exhibitions": [
            {{
            "title": "...",
            "venue": "...",
            "is_artistic_showcase": true | false,
            "is_virtual": true | false,
            "authorship_verified": true | false,
            "evidence_summary": ["..."],
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        response = RFEBOT.reponseParser(llm2(prompt))
        return response  # this is the reponse for the criterion 8

    @staticmethod
    def AgenticCriterionPrompt8(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 8**:  **"Evidence that the beneficiary has performed in a leading or critical role for organizations or establishments that have a distinguished reputation."**

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Instructions:

        1. Evaluate each organization where the petitioner claims to have served in a **leading or critical role**.
        - “Leading” means high-ranking decision-making position (e.g., director, CTO, head researcher)
        - “Critical” means the work was essential to the outcome or success of the organization’s goals

        2. For each organization:
        - `name`: Organization's name
        - `beneficiary_role`: Stated role or position held by the petitioner
        - `is_leading_or_critical`: `true` if the evidence shows the petitioner's performance significantly influenced the organization
        - `role_explanation`: Describe what specific actions or performance justifies the role as leading/critical
        - `distinguished_reputation`: `true` if the organization is shown to be of high repute in its industry/field
        - `reputation_evidence`: List sources (e.g., media recognition, awards, rankings, partnerships, market share, historical significance, etc.)
        - `supporting_documents`: List submitted materials such as:
            - Organizational charts
            - Letters from CEOs or stakeholders
            - Internal reports, KPIs influenced by petitioner
            - Press releases or awards
            - Public news/media coverage of the petitioner's impact or the organization's prominence
        - `issues`: Highlight any problems, e.g., unclear responsibilities, weak reputation evidence, unclear link between role and outcome

        3. `meets_criterion_8: true` only if **at least one** organization qualifies under:
        - Role is well-evidenced as leading or critical
        - Organization is demonstrated to be of **distinguished reputation**

        4. Use `missing_elements` to list gaps in the evidence, such as:
        - Lack of performance details in the role
        - No documents proving the organization's prominence
        - Role is high in title but low in impact

        5. Use `suggested_supporting_evidence` to recommend how to improve, such as:
        - Submit detailed letters explaining specific impact of the role
        - Provide company press coverage, public financial data, industry ranking
        - Clarify the unique significance of the ' contribution vs. others

        ALWAYS Return the result in the following JSON format. Do not include any commentary or explanations outside the structure.
        {{
        "meets_criterion_8": true | false,
        "organizations": [
            {{
            "name": "...",
            "beneficiary_role": "...",
            "is_leading_or_critical": true | false,
            "role_explanation": "...",
            "distinguished_reputation": true | false,
            "reputation_evidence": ["..."],
            "supporting_documents": ["..."],
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        response = RFEBOT.reponseParser(llm2(prompt))
        return response  # this is the response for the criterion 8

    @staticmethod
    def AgenticCriterionPrompt9(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 9**:  **"Evidence that the beneficiary has commanded a high salary or other significantly high remuneration for services, in relation to others in the field."**

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Instructions:

        1. Evaluate whether the petition demonstrates that the beneficiary received a **high salary or other significantly high remuneration**, and whether it is **high relative to others in the field**.

        2. For each salary/remuneration claim:
        - `year`: The tax year or employment period the evidence covers
        - `salary_amount`: Total salary/remuneration stated (with currency)
        - `source`: Type of evidence (e.g., W-2, 1099, foreign tax form, salary letter, contract)
        - `comparison_evidence_provided`: `true` if the petition includes documents showing how this salary compares to others in the same field/region
        - `is_significantly_high`: `true` if supported evidence clearly shows the salary is above industry norms
        - `supporting_documents`: List submitted materials such as:
            - Pay stubs, W-2s, 1099s, foreign tax records
            - Salary letters or employment contracts
            - Compensation surveys, public salary databases
            - Media reports, awards tied to high compensation
        - `issues`: Highlight any concerns (e.g., no comparison data, role unrelated to field, unclear if base salary or total comp)

        3. `meets_criterion_9: true` only if at least one salary is:
        - Shown to be **significantly high relative** to others in the **same or allied field**
        - Supported by comparative industry data or documentation

        4. Use `missing_elements` to list gaps in the evidence, such as:
        - Lack of comparative compensation data
        - No supporting documents (just a statement or letter)
        - Unclear if the remuneration is for work in the claimed field of expertise

        5. Use `suggested_supporting_evidence` to recommend how to improve, such as:
        - Provide industry-wide salary comparison charts
        - Submit compensation data from DOL, Glassdoor, PayScale, etc.
        - Add letters from employers justifying above-market compensation

        ALWAYS Return the result in the following JSON format. Do not include any commentary or explanations outside the structure.

        {{
        "meets_criterion_9": true | false,
        "salary_records": [
            {{
            "year": "...",
            "salary_amount": "...",
            "source": "...",
            "comparison_evidence_provided": true | false,
            "is_significantly_high": true | false,
            "supporting_documents": ["..."],
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        response = RFEBOT.reponseParser(llm2(prompt))
        return response  # this is the reponse for criterion 9

    @staticmethod
    def AgenticCriterionPrompt10(petitionContent: str) -> Dict:
        prompt: str = f"""You are a USCIS Officer reviewing an EB-1A petition. Your task is to evaluate whether the petitioner satisfies **Criterion 10**:  **"Evidence of commercial successes in the performing arts, as shown by box office receipts or record, cassette, compact disk, or video sales."**

        <petitionContent>
        {petitionContent.strip()}
        </petitionContent>

        Instructions:

        1. Determine whether the beneficiary has demonstrated **commercial success in the performing arts**.
        - Commercial success must be evidenced by **volume of sales** or **box office receipts**
        - Merely performing or releasing work does **not** prove commercial success

        2. For each production or work claimed:
        - `title`: Name of the work (e.g., movie, album, production)
        - `type`: Type of performing art (e.g., music, theater, film, TV)
        - `sales_or_receipts`: Total reported sales or earnings (include currency and figures)
        - `is_commercial_success`: `true` if sales/receipts are substantial compared to industry benchmarks
        - `supporting_documents`: List submitted materials such as:
            - Box office statements
            - Record or CD/DVD sales reports
            - Streaming platform revenue stats
            - Contracts with financial breakdowns
            - Third-party evidence showing comparative success
        - `issues`: Any concerns (e.g., no sales numbers, anecdotal success, unrelated performances)

        3. `meets_criterion_10: true` only if at least one work is:
        - **Commercially successful** based on volume of sales/receipts
        - Supported with concrete, quantifiable evidence

        4. Use `missing_elements` to list gaps in the evidence, such as:
        - No proof of actual sales or earnings
        - No comparative data to show that success is significant
        - Just mentions of performances without financial metrics

        5. Use `suggested_supporting_evidence` to recommend improvements, such as:
        - Submit box office or ticketing statements with verified totals
        - Provide music or video sales reports (e.g., Nielsen, Spotify, Billboard, iTunes)
        - Include industry comparisons showing top percentile revenue

        ALWAYS Return the result in the following JSON format. Do not include any commentary or explanations outside the structure.

        {{
        "meets_criterion_10": true | false,
        "works": [
            {{
            "title": "...",
            "type": "...",
            "sales_or_receipts": "...",
            "is_commercial_success": true | false,
            "supporting_documents": ["..."],
            "issues": ["..."]
            }}
        ],
        "overall_reasoning": "...",
        "missing_elements": ["..."],
        "suggested_supporting_evidence": ["..."]
        }}
        """
        response = RFEBOT.reponseParser(llm2(prompt))
        return response  # this is the reponse for the criterion 10
    
    @staticmethod
    def finalAgent(result: dict) -> str:
        print("Exporting the Analysis into a Doc File")
        doc = Document()
        doc.add_heading("Criterion EB-1A Petition Assessment Summary", level=1)

        recognizable_components = result.get("recognizable_components", {})
        background = recognizable_components.get("personal_background", "")
        expert_letters = recognizable_components.get(
            "expert_recommendation_letters", []
        )
        media_and_judging = recognizable_components.get("media_and_judging_roles", [])
        red_flags = recognizable_components.get("red_flags", {})
        evidence = recognizable_components.get("evidence", {})

        doc.add_heading("1. Personal Background", level=2)
        doc.add_paragraph(background)

        doc.add_heading("2. Expert Recommendation Letters", level=2)
        for letter in expert_letters:
            doc.add_paragraph(f" {letter}")

        doc.add_heading("3. Media and Judging Roles", level=2)
        for role in media_and_judging:
            doc.add_paragraph(f" {role}")

        doc.add_heading("4. Red Flags", level=2)
        for k, v in red_flags.items():
            title = k.replace("_", " ").title()
            doc.add_paragraph(f"{title}:\n{v}", style="List Bullet")

        doc.add_heading("5. Criteria-specific evidence", level=2)
        if "lesser_awards" in evidence:
            for item in evidence["lesser_awards"]:
                doc.add_paragraph(f"{item}", style="List Bullet")

        if "association_membership" in evidence:
            for item in evidence["association_membership"]:
                doc.add_paragraph(f"{item}", style="List Bullet")

        if "published_material" in evidence:
            for item in evidence["published_material"]:
                doc.add_paragraph(f"{item}", style="List Bullet")

        if "judging_work_of_others" in evidence:
            for item in evidence["judging_work_of_others"]:
                doc.add_paragraph(f"{item}", style="List Bullet")

        if "original_contributions" in evidence:
            for item in evidence["original_contributions"]:
                doc.add_paragraph(f"{item}", style="List Bullet")


        if "scholarly_articles" in evidence:
            for item in evidence["scholarly_articles"]:
                doc.add_paragraph(f"{item}", style="List Bullet")
        

        if "artistic_display" in evidence:
            for item in evidence["artistic_display"]:
                doc.add_paragraph(f"{item}", style="List Bullet")


        if "leading_or_critical_role" in evidence:
            for item in evidence["leading_or_critical_role"]:
                doc.add_paragraph(f"{item}", style="List Bullet")


        if "high_salary_or_remuneration" in evidence:
            for item in evidence["high_salary_or_remuneration"]:
                doc.add_paragraph(f"{item}", style="List Bullet")


        if "commercial_success_in_arts" in evidence:
            for item in evidence["jcommercial_success_in_arts"]:
                doc.add_paragraph(f"{item}", style="List Bullet")


        summary = result.get("summary", {})
        doc.add_heading("5. Summary", level=2)
        doc.add_paragraph(
            f"Qualification Path: {summary.get('qualification_path', '')}"
        )
        doc.add_paragraph(f"Status: {summary.get('status', '')}")
        doc.add_paragraph(f"Criteria Met: {summary.get('criteria_met', '')}")
        doc.add_paragraph(
            f"Total Criteria Checked: {summary.get('total_criteria_checked', '')}"
        )

        criterionTitles = {
            "Criterion1Result": "Criterion 1: Documentation of the alien's receipt of lesser nationally or internationally recognized prizes or awards for excellence in the field of endeavor.",
            "Criterion2Result": "Criterion 2: Documentation of the alien's membership in associations in the field which demand outstanding achievement of their members.",
            "Criterion3Result": "Criterion 3: Published material about the alien in professional or major trade publications or other major media, relating to the alien's work in the field.",
            "Criterion4Result": "Criterion 4: Evidence of the alien's participation, either individually or on a panel, as a judge of the work of others in the same or an allied field.",
            "Criterion5Result": "Criterion 5: Evidence of the alien's original scientific, scholarly, artistic, athletic, or business-related contributions of major significance in the field.",
            "Criterion6Result": "Criterion 6: Evidence of the alien's authorship of scholarly articles in professional or major trade publications or other major media.",
            "Criterion7Result": "Criterion 7: Evidence of the display of the alien's work in exhibitions or showcases.",
            "Criterion8Result": "Criterion 8: Evidence that the alien has performed in a leading or critical role for organizations or establishments that have a distinguished reputation.",
            "Criterion9Result": "Criterion 9: Evidence that the alien has commanded a high salary or other significantly high remuneration for services in relation to others in the field.",
            "Criterion10Result": "Criterion 10: Evidence of commercial successes in the performing arts, as shown by box office receipts or record, cassette, compact disk, or video sales.",
        }

        doc.add_heading("6. Detailed Criterion Evaluation", level=2)

        for i in range(1, 11):
            criterion_key = f"Criterion{i}Result"
            criterion_data = result.get(criterion_key, {})
            doc.add_heading(
                criterionTitles.get(criterion_key, f"Criterion {i}"), level=3
            )

            if not criterion_data:
                doc.add_paragraph("No data provided.")
                continue

            if criterion_data.get("status") == "couldn't parse the content":
                doc.add_paragraph("Could not parse this criterion.")
                continue

            doc.add_paragraph(
                f"Meets Criterion: {criterion_data.get(f'meets_criterion_{i}', 'Unknown')}"
            )

            reasoning = criterion_data.get("reasoning") or criterion_data.get(
                "overall_reasoning", ""
            )
            if reasoning:
                doc.add_paragraph(f"Reasoning:\n{reasoning}")

            if "issues_detected" in criterion_data:
                doc.add_paragraph("Issues Detected:")
                for issue in criterion_data["issues_detected"]:
                    doc.add_paragraph(f" {issue}", style="List Bullet")

            if "award_names" in criterion_data:
                doc.add_paragraph("Awards:")
                for item in criterion_data["award_names"]:
                    doc.add_paragraph(f" {item}", style="List Bullet")

            if "association_names" in criterion_data:
                doc.add_paragraph("Associations:")
                for item in criterion_data["association_names"]:
                    doc.add_paragraph(f" {item}", style="List Bullet")

            if "contributions" in criterion_data:
                doc.add_paragraph("Contributions:")
                for contrib in criterion_data["contributions"]:
                    doc.add_paragraph(
                        f" {contrib.get('title_or_description', '')}",
                        style="List Bullet",
                    )
                    impact = contrib.get("impact_summary", "")
                    if impact:
                        doc.add_paragraph(f"  Impact: {impact}")

            if "salary_records" in criterion_data:
                doc.add_paragraph("Salary Records:")
                for salary in criterion_data["salary_records"]:
                    amount = salary.get("salary_amount", "")
                    year = salary.get("year", "")
                    doc.add_paragraph(f" {year}: {amount}", style="List Bullet")

            if "articles" in criterion_data:
                doc.add_paragraph("Articles:")
                for art in criterion_data["articles"]:
                    doc.add_paragraph(
                        f" {art.get('title', '')} ({art.get('publication', '')})",
                        style="List Bullet",
                    )

            if "judging_instances" in criterion_data:
                doc.add_paragraph("Judging Roles:")
                for judge in criterion_data["judging_instances"]:
                    doc.add_paragraph(
                        f" {judge.get('event_name_or_platform', '')}: {judge.get('role', '')}",
                        style="List Bullet",
                    )

            if "missing_elements" in criterion_data:
                doc.add_paragraph("Missing Elements:")
                for missing in criterion_data["missing_elements"]:
                    doc.add_paragraph(f" {missing}", style="List Bullet")

            if "suggested_supporting_evidence" in criterion_data:
                doc.add_paragraph("Suggested Supporting Evidence:")
                for evidence in criterion_data["suggested_supporting_evidence"]:
                    doc.add_paragraph(f" {evidence}", style="List Bullet")

        output_path = "Criterion_EB1A_Assessment_Summary.docx"
        doc.save(output_path)
        return output_path
